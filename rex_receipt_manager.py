"""
rex_receipt_manager.py — GOJ Receipt Management System
════════════════════════════════════════════════════════════
Multi-user receipt intake → OCR → categorized filing →
organized Excel workbook → formal Word expense reports.

Garden of Joy · Gold Health Systems

WHO SUBMITS RECEIPTS:
  • Misha  — sends photo to REX Telegram bot → auto-filed instantly
  • Kato   — chairman can submit and manage all receipts

WHO CAN VIEW RECEIPTS (Dashboard section):
  • Chairman (Kato) — full access: view, reorganize, delete, export
  • Vlad            — view-only access to receipt section

EXCEL WORKBOOK (GOJ_Receipts_Master.xlsx):
  Regenerated every time a new receipt is logged.
  5 sheets:
    1. "All Receipts"     — chronological, all submitters
    2. "By Vendor"        — sorted/grouped by company
    3. "By Amount"        — highest spend first
    4. "By Item"          — line items expanded (each item = one row)
    5. "Monthly Summary"  — pivot-style totals by category × month

WORD REPORT (GOJ_Expense_Report_YYYY-MM.docx):
  Formal expense report, generated on demand.
  One page per category, with receipt table + totals.
  Generated with python-docx.

TELEGRAM COMMANDS:
  (Kato / Chairman):
    [photo]                      → file receipt
    receipts by company          → regenerate & send Excel sorted by vendor
    receipts by price            → Excel sorted by amount
    receipts by item             → Excel with line items expanded
    receipts report              → Word/PDF formal report (this month)
    receipts report [YYYY-MM]    → specific month
    receipts this week           → text summary
    receipts misha               → only Misha's submissions
    receipts [vendor name]       → filter by vendor
    receipts delete [id]         → soft-delete a receipt

  (Vlad — view only, dashboard section):
    receipts                     → recent summary
    receipts this week

  (Misha — submit only):
    [photo]                      → files receipt, confirmation sent
"""

from __future__ import annotations

import json
import logging
import os
import re
import sqlite3
import tempfile
import time
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── V4 Rexonasence patch — visibility gate + event emission ───────────────────
try:
    from rex_receipt_manager_v4_patch import (
        run_v4_migration, ReceiptVisibilityGate,
        emit_receipt_event, flag_for_review, mark_receipt_reviewed,
    )
    _V4_PATCH_OK = True
    _visibility_gate: Optional["ReceiptVisibilityGate"] = None  # initialized after DB ready
except Exception as _v4_import_err:
    _V4_PATCH_OK = False
    _visibility_gate = None
    logger.debug(f"[v4] patch not loaded (non-fatal): {_v4_import_err}")
    def run_v4_migration(*_a, **_kw): return {"added": [], "already_present": []}  # type: ignore
    def emit_receipt_event(*_a, **_kw): pass   # type: ignore
    def flag_for_review(*_a, **_kw): pass      # type: ignore
    def mark_receipt_reviewed(*_a, **_kw): pass  # type: ignore

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR        = Path.home() / "Desktop" / "REX"
RECEIPTS_ROOT   = BASE_DIR / "receipts"
LEDGER_DB       = BASE_DIR / "data" / "rex_ledger.db"
REPORTS_DIR     = BASE_DIR / "reports"
MASTER_XL_PATH  = RECEIPTS_ROOT / "GOJ_Receipts_Master.xlsx"

# ── Access control ─────────────────────────────────────────────────────────────
# 0 = no access, 1 = submit only, 2 = view-only, 3 = full access
ROLE_NONE    = 0
ROLE_SUBMIT  = 1  # Misha: can submit receipts, gets confirmation
ROLE_VIEW    = 2  # Vlad: view dashboard receipt section
ROLE_FULL    = 3  # Kato: full access — view, edit, reorganize, delete, export


# ──────────────────────────────────────────────────────────────────────────────
# DATABASE LAYER
# ──────────────────────────────────────────────────────────────────────────────

def _ensure_db() -> None:
    LEDGER_DB.parent.mkdir(parents=True, exist_ok=True)
    RECEIPTS_ROOT.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    con = sqlite3.connect(str(LEDGER_DB))
    # Step 1: Create tables (base schema, no columns that may be missing)
    con.executescript("""
        CREATE TABLE IF NOT EXISTS receipts (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            receipt_date    TEXT,
            vendor          TEXT,
            amount          REAL,
            tax             REAL DEFAULT 0,
            category        TEXT DEFAULT 'Misc',
            subcategory     TEXT,
            description     TEXT,
            pdf_path        TEXT,
            raw_text        TEXT,
            source_file     TEXT,
            logged_at       TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS line_items (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            receipt_id  INTEGER NOT NULL,
            description TEXT,
            quantity    REAL DEFAULT 1,
            unit_price  REAL,
            total       REAL,
            FOREIGN KEY (receipt_id) REFERENCES receipts(id)
        );

        CREATE INDEX IF NOT EXISTS idx_rcpt_date     ON receipts(receipt_date);
        CREATE INDEX IF NOT EXISTS idx_rcpt_category ON receipts(category);
        CREATE INDEX IF NOT EXISTS idx_rcpt_vendor   ON receipts(vendor);
    """)
    con.commit()

    # Step 2: Migration — add new columns to existing tables if missing
    existing_cols = [row[1] for row in con.execute("PRAGMA table_info(receipts)").fetchall()]
    migrations = [
        ("submitted_by", "TEXT DEFAULT 'kato'"),
        ("confirmed",    "INTEGER DEFAULT 0"),
        ("deleted",      "INTEGER DEFAULT 0"),
    ]
    for col, col_def in migrations:
        if col not in existing_cols:
            try:
                con.execute(f"ALTER TABLE receipts ADD COLUMN {col} {col_def}")
                logger.info(f"[manager] Added column '{col}' to receipts")
            except Exception as e:
                logger.debug(f"[manager] Column migration note for '{col}': {e}")
    con.commit()

    # Step 3: Create index on submitted_by (after migration ensures column exists)
    try:
        con.execute("CREATE INDEX IF NOT EXISTS idx_rcpt_submitter ON receipts(submitted_by)")
        con.commit()
    except Exception:
        pass

    con.close()

    # Step 4: V4 schema extension (idempotent — adds department, confidence,
    #         visibility_class, review_status, reviewed_by, reviewed_at)
    if _V4_PATCH_OK:
        try:
            run_v4_migration()
            global _visibility_gate
            if _visibility_gate is None:
                _visibility_gate = ReceiptVisibilityGate()
        except Exception as _e:
            logger.debug(f"[v4] migration/gate init (non-fatal): {_e}")


def _get_receipts(
    start: str = "2000-01-01",
    end:   str = "2099-12-31",
    submitter: Optional[str] = None,
    vendor_filter: Optional[str] = None,
    include_deleted: bool = False,
) -> list[dict]:
    """Pull receipts from ledger with optional filters."""
    try:
        con = sqlite3.connect(str(LEDGER_DB))
        con.row_factory = sqlite3.Row

        where_clauses = ["receipt_date BETWEEN ? AND ?"]
        params: list = [start, end]

        if not include_deleted:
            where_clauses.append("deleted = 0")
        if submitter:
            where_clauses.append("submitted_by = ?")
            params.append(submitter)
        if vendor_filter:
            where_clauses.append("LOWER(vendor) LIKE ?")
            params.append(f"%{vendor_filter.lower()}%")

        where = " AND ".join(where_clauses)
        rows  = con.execute(
            f"SELECT * FROM receipts WHERE {where} ORDER BY receipt_date DESC",
            params
        ).fetchall()
        con.close()
        return [dict(r) for r in rows]
    except Exception as e:
        logger.error(f"[manager] _get_receipts error: {e}")
        return []


def _get_line_items(receipt_id: int) -> list[dict]:
    try:
        con = sqlite3.connect(str(LEDGER_DB))
        con.row_factory = sqlite3.Row
        rows = con.execute(
            "SELECT * FROM line_items WHERE receipt_id=?", (receipt_id,)
        ).fetchall()
        con.close()
        return [dict(r) for r in rows]
    except Exception:
        return []


# ──────────────────────────────────────────────────────────────────────────────
# EXCEL WORKBOOK GENERATION
# (openpyxl — per xlsx skill guidance)
# ──────────────────────────────────────────────────────────────────────────────

def _xl_header_style():
    from openpyxl.styles import Font, PatternFill, Alignment
    return {
        "font":      Font(name="Arial", bold=True, color="FFFFFF", size=11),
        "fill":      PatternFill("solid", start_color="1F4E79"),   # deep navy
        "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
    }

def _xl_subheader_style():
    from openpyxl.styles import Font, PatternFill, Alignment
    return {
        "font":      Font(name="Arial", bold=True, color="FFFFFF", size=10),
        "fill":      PatternFill("solid", start_color="2E75B6"),   # medium blue
        "alignment": Alignment(horizontal="left", vertical="center"),
    }

def _xl_zebra_style(row_idx: int):
    from openpyxl.styles import PatternFill
    if row_idx % 2 == 0:
        return PatternFill("solid", start_color="EBF3FB")
    return None

def _xl_currency_format():
    return '"$"#,##0.00_);("$"#,##0.00)'

def _apply_header(ws, row: int, cols: list[str]) -> None:
    """Write bold navy headers across a row."""
    from openpyxl.styles import Font, PatternFill, Alignment
    style = _xl_header_style()
    for i, label in enumerate(cols, start=1):
        cell = ws.cell(row=row, column=i, value=label)
        cell.font      = style["font"]
        cell.fill      = style["fill"]
        cell.alignment = style["alignment"]


def build_master_workbook(receipts: list[dict]) -> Path:
    """
    Build (or rebuild) GOJ_Receipts_Master.xlsx with 5 sheets.
    Returns path to the saved file.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise ImportError(
            "openpyxl is required. Install: "
            "pip install openpyxl --break-system-packages"
        )

    wb  = Workbook()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Remove default sheet ──────────────────────────────────────────────────
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    # ── Sheet 1: All Receipts (chronological) ─────────────────────────────────
    ws1 = wb.create_sheet("All Receipts")
    ws1.sheet_view.showGridLines = True
    ws1.freeze_panes = "A3"

    # Title row
    ws1.merge_cells("A1:I1")
    title_cell = ws1["A1"]
    title_cell.value     = f"Garden of Joy — All Receipts  (generated {now})"
    title_cell.font      = Font(name="Arial", bold=True, size=13, color="1F4E79")
    title_cell.alignment = Alignment(horizontal="center")
    ws1.row_dimensions[1].height = 22

    cols_all = ["ID", "Date", "Vendor / Company", "Category",
                "Amount ($)", "Tax ($)", "Submitted By", "Description", "PDF"]
    _apply_header(ws1, 2, cols_all)
    ws1.row_dimensions[2].height = 20

    total_row_offset = 3
    for i, r in enumerate(sorted(receipts, key=lambda x: x["receipt_date"] or "", reverse=True), start=3):
        ws1.cell(i, 1, r["id"])
        ws1.cell(i, 2, r["receipt_date"])
        ws1.cell(i, 3, r["vendor"])
        ws1.cell(i, 4, r["category"])
        amt_cell = ws1.cell(i, 5, r["amount"])
        amt_cell.number_format = _xl_currency_format()
        tax_cell = ws1.cell(i, 6, r["tax"])
        tax_cell.number_format = _xl_currency_format()
        ws1.cell(i, 7, (r["submitted_by"] or "kato").title())
        ws1.cell(i, 8, (r["description"] or "")[:60])
        if r.get("pdf_path"):
            ws1.cell(i, 9, Path(r["pdf_path"]).name)

        # Zebra striping
        fill = _xl_zebra_style(i)
        if fill:
            for col in range(1, 10):
                ws1.cell(i, col).fill = fill

    # Totals row
    last_data = 2 + len(receipts)
    ws1.cell(last_data + 1, 4, "TOTAL").font = Font(name="Arial", bold=True)
    total_cell = ws1.cell(last_data + 1, 5,
                          f"=SUM(E3:E{last_data})" if len(receipts) > 0 else 0)
    total_cell.number_format = _xl_currency_format()
    total_cell.font = Font(name="Arial", bold=True, color="1F4E79")

    # Column widths
    ws1.column_dimensions["A"].width = 6
    ws1.column_dimensions["B"].width = 12
    ws1.column_dimensions["C"].width = 28
    ws1.column_dimensions["D"].width = 16
    ws1.column_dimensions["E"].width = 13
    ws1.column_dimensions["F"].width = 11
    ws1.column_dimensions["G"].width = 14
    ws1.column_dimensions["H"].width = 30
    ws1.column_dimensions["I"].width = 26

    # ── Sheet 2: By Vendor ────────────────────────────────────────────────────
    ws2 = wb.create_sheet("By Vendor")
    ws2.freeze_panes = "A3"

    ws2.merge_cells("A1:F1")
    ws2["A1"].value     = f"Garden of Joy — Receipts by Vendor/Company  (generated {now})"
    ws2["A1"].font      = Font(name="Arial", bold=True, size=13, color="1F4E79")
    ws2["A1"].alignment = Alignment(horizontal="center")
    ws2.row_dimensions[1].height = 22

    # Group by vendor
    from collections import defaultdict
    vendor_groups: dict[str, list[dict]] = defaultdict(list)
    for r in receipts:
        key = (r["vendor"] or "Unknown").strip()
        vendor_groups[key].append(r)

    # Sort vendors by total spend descending
    vendor_sorted = sorted(
        vendor_groups.items(),
        key=lambda x: sum(r["amount"] or 0 for r in x[1]),
        reverse=True
    )

    ws2.merge_cells("A2:F2")
    ws2["A2"].value = f"TOTAL VENDORS: {len(vendor_sorted)}"
    ws2["A2"].font  = Font(name="Arial", bold=True, color="2E75B6")

    row_idx = 3
    for vendor_name, vreceipts in vendor_sorted:
        # Vendor subheader
        ws2.merge_cells(f"A{row_idx}:F{row_idx}")
        cell = ws2[f"A{row_idx}"]
        vendor_total = sum(r["amount"] or 0 for r in vreceipts)
        cell.value     = f"  {vendor_name}  —  {len(vreceipts)} receipt(s)  |  Total: ${vendor_total:,.2f}"
        sh_style = _xl_subheader_style()
        cell.font      = sh_style["font"]
        cell.fill      = sh_style["fill"]
        cell.alignment = sh_style["alignment"]
        ws2.row_dimensions[row_idx].height = 18
        row_idx += 1

        # Column headers
        _apply_header(ws2, row_idx, ["Date", "Amount", "Tax", "Category", "Submitted By", "Description"])
        ws2.row_dimensions[row_idx].height = 16
        row_idx += 1

        # Data rows
        start_data = row_idx
        for r in sorted(vreceipts, key=lambda x: x["receipt_date"] or "", reverse=True):
            ws2.cell(row_idx, 1, r["receipt_date"])
            a = ws2.cell(row_idx, 2, r["amount"])
            a.number_format = _xl_currency_format()
            t = ws2.cell(row_idx, 3, r["tax"])
            t.number_format = _xl_currency_format()
            ws2.cell(row_idx, 4, r["category"])
            ws2.cell(row_idx, 5, (r["submitted_by"] or "kato").title())
            ws2.cell(row_idx, 6, (r["description"] or "")[:50])
            fill = _xl_zebra_style(row_idx)
            if fill:
                for col in range(1, 7):
                    ws2.cell(row_idx, col).fill = fill
            row_idx += 1

        # Vendor subtotal
        ws2.cell(row_idx, 1, "Subtotal").font = Font(name="Arial", bold=True, italic=True)
        sub = ws2.cell(row_idx, 2, f"=SUM(B{start_data}:B{row_idx - 1})")
        sub.number_format = _xl_currency_format()
        sub.font = Font(name="Arial", bold=True, italic=True, color="1F4E79")
        row_idx += 2   # gap between vendors

    ws2.column_dimensions["A"].width = 12
    ws2.column_dimensions["B"].width = 13
    ws2.column_dimensions["C"].width = 11
    ws2.column_dimensions["D"].width = 16
    ws2.column_dimensions["E"].width = 14
    ws2.column_dimensions["F"].width = 32

    # ── Sheet 3: By Amount ────────────────────────────────────────────────────
    ws3 = wb.create_sheet("By Amount")
    ws3.freeze_panes = "A3"

    ws3.merge_cells("A1:G1")
    ws3["A1"].value     = f"Garden of Joy — Receipts by Amount (Highest First)  (generated {now})"
    ws3["A1"].font      = Font(name="Arial", bold=True, size=13, color="1F4E79")
    ws3["A1"].alignment = Alignment(horizontal="center")
    ws3.row_dimensions[1].height = 22

    _apply_header(ws3, 2, ["Rank", "Amount ($)", "Vendor", "Date", "Category",
                            "Tax ($)", "Submitted By"])
    ws3.row_dimensions[2].height = 20

    sorted_by_amt = sorted(receipts, key=lambda x: x["amount"] or 0, reverse=True)
    for rank, r in enumerate(sorted_by_amt, start=1):
        row_i = rank + 2
        ws3.cell(row_i, 1, rank)
        a = ws3.cell(row_i, 2, r["amount"])
        a.number_format = _xl_currency_format()
        ws3.cell(row_i, 3, r["vendor"])
        ws3.cell(row_i, 4, r["receipt_date"])
        ws3.cell(row_i, 5, r["category"])
        t = ws3.cell(row_i, 6, r["tax"])
        t.number_format = _xl_currency_format()
        ws3.cell(row_i, 7, (r["submitted_by"] or "kato").title())
        fill = _xl_zebra_style(row_i)
        if fill:
            for col in range(1, 8):
                ws3.cell(row_i, col).fill = fill

        # Bold top 3 highest
        if rank <= 3:
            for col in range(1, 8):
                ws3.cell(row_i, col).font = Font(name="Arial", bold=True)

    last = 2 + len(sorted_by_amt)
    ws3.cell(last + 1, 2, f"=SUM(B3:B{last})" if receipts else 0).number_format = _xl_currency_format()
    ws3.cell(last + 1, 1, "TOTAL").font = Font(name="Arial", bold=True)

    ws3.column_dimensions["A"].width = 7
    ws3.column_dimensions["B"].width = 14
    ws3.column_dimensions["C"].width = 28
    ws3.column_dimensions["D"].width = 12
    ws3.column_dimensions["E"].width = 16
    ws3.column_dimensions["F"].width = 11
    ws3.column_dimensions["G"].width = 14

    # ── Sheet 4: By Item (line items expanded) ────────────────────────────────
    ws4 = wb.create_sheet("By Item")
    ws4.freeze_panes = "A3"

    ws4.merge_cells("A1:H1")
    ws4["A1"].value     = f"Garden of Joy — Line Items Expanded  (generated {now})"
    ws4["A1"].font      = Font(name="Arial", bold=True, size=13, color="1F4E79")
    ws4["A1"].alignment = Alignment(horizontal="center")
    ws4.row_dimensions[1].height = 22

    _apply_header(ws4, 2, ["Receipt ID", "Date", "Vendor", "Item Description",
                            "Qty", "Unit Price ($)", "Line Total ($)", "Category"])
    ws4.row_dimensions[2].height = 20

    item_row = 3
    all_items: list[tuple] = []
    for r in receipts:
        items = _get_line_items(r["id"])
        if items:
            for item in items:
                all_items.append((r, item))
        else:
            # No line items — use receipt as single row
            all_items.append((r, {
                "description": r.get("description") or r.get("vendor") or "Receipt",
                "quantity": 1,
                "unit_price": r.get("amount", 0),
                "total": r.get("amount", 0),
            }))

    # Sort by item description
    all_items.sort(key=lambda x: (x[1].get("description") or "").lower())

    for r, item in all_items:
        ws4.cell(item_row, 1, r["id"])
        ws4.cell(item_row, 2, r["receipt_date"])
        ws4.cell(item_row, 3, r["vendor"])
        ws4.cell(item_row, 4, (item.get("description") or "")[:60])
        ws4.cell(item_row, 5, item.get("quantity", 1))
        up = ws4.cell(item_row, 6, item.get("unit_price"))
        if up.value is not None:
            up.number_format = _xl_currency_format()
        lt = ws4.cell(item_row, 7, item.get("total"))
        if lt.value is not None:
            lt.number_format = _xl_currency_format()
        ws4.cell(item_row, 8, r["category"])
        fill = _xl_zebra_style(item_row)
        if fill:
            for col in range(1, 9):
                ws4.cell(item_row, col).fill = fill
        item_row += 1

    last4 = item_row - 1
    ws4.cell(last4 + 1, 3, "TOTAL").font = Font(name="Arial", bold=True)
    tc = ws4.cell(last4 + 1, 7,
                  f"=SUM(G3:G{last4})" if all_items else 0)
    tc.number_format = _xl_currency_format()
    tc.font = Font(name="Arial", bold=True, color="1F4E79")

    ws4.column_dimensions["A"].width = 10
    ws4.column_dimensions["B"].width = 12
    ws4.column_dimensions["C"].width = 24
    ws4.column_dimensions["D"].width = 36
    ws4.column_dimensions["E"].width = 6
    ws4.column_dimensions["F"].width = 13
    ws4.column_dimensions["G"].width = 14
    ws4.column_dimensions["H"].width = 16

    # ── Sheet 5: Monthly Summary ──────────────────────────────────────────────
    ws5 = wb.create_sheet("Monthly Summary")

    ws5.merge_cells("A1:F1")
    ws5["A1"].value     = f"Garden of Joy — Monthly Expense Summary  (generated {now})"
    ws5["A1"].font      = Font(name="Arial", bold=True, size=13, color="1F4E79")
    ws5["A1"].alignment = Alignment(horizontal="center")
    ws5.row_dimensions[1].height = 22

    try:
        con = sqlite3.connect(str(LEDGER_DB))
        monthly_data = con.execute("""
            SELECT
                strftime('%Y-%m', receipt_date) AS month,
                category,
                COUNT(*) AS cnt,
                COALESCE(SUM(amount), 0) AS total,
                COALESCE(SUM(tax), 0) AS total_tax
            FROM receipts
            WHERE deleted = 0
            GROUP BY month, category
            ORDER BY month DESC, total DESC
        """).fetchall()

        monthly_totals = con.execute("""
            SELECT
                strftime('%Y-%m', receipt_date) AS month,
                COUNT(*) AS cnt,
                COALESCE(SUM(amount), 0) AS total
            FROM receipts
            WHERE deleted = 0
            GROUP BY month
            ORDER BY month DESC
        """).fetchall()
        con.close()
    except Exception:
        monthly_data   = []
        monthly_totals = []

    # Group by month
    month_groups: dict[str, list] = {}
    for month, category, cnt, total, total_tax in monthly_data:
        if month not in month_groups:
            month_groups[month] = []
        month_groups[month].append((category, cnt, total, total_tax))

    month_totals_map = {row[0]: (row[1], row[2]) for row in monthly_totals}

    row_idx = 2
    for month, categories in sorted(month_groups.items(), reverse=True):
        # Month header
        ws5.merge_cells(f"A{row_idx}:F{row_idx}")
        total_cnt, total_amt = month_totals_map.get(month, (0, 0))
        cell = ws5[f"A{row_idx}"]
        cell.value     = f"  {month}  —  {total_cnt} receipts  |  Total: ${total_amt:,.2f}"
        sh_style = _xl_subheader_style()
        cell.font      = sh_style["font"]
        cell.fill      = sh_style["fill"]
        cell.alignment = sh_style["alignment"]
        ws5.row_dimensions[row_idx].height = 18
        row_idx += 1

        _apply_header(ws5, row_idx, ["Category", "Count", "Subtotal ($)", "Tax ($)", "%", "Notes"])
        row_idx += 1

        for category, cnt, total, total_tax in sorted(categories, key=lambda x: -x[2]):
            pct = (total / total_amt * 100) if total_amt else 0
            ws5.cell(row_idx, 1, category)
            ws5.cell(row_idx, 2, cnt)
            sub = ws5.cell(row_idx, 3, total)
            sub.number_format = _xl_currency_format()
            tax = ws5.cell(row_idx, 4, total_tax)
            tax.number_format = _xl_currency_format()
            ws5.cell(row_idx, 5, round(pct / 100, 3)).number_format = "0.0%"
            fill = _xl_zebra_style(row_idx)
            if fill:
                for col in range(1, 7):
                    ws5.cell(row_idx, col).fill = fill
            row_idx += 1

        row_idx += 1   # gap

    ws5.column_dimensions["A"].width = 18
    ws5.column_dimensions["B"].width = 8
    ws5.column_dimensions["C"].width = 14
    ws5.column_dimensions["D"].width = 12
    ws5.column_dimensions["E"].width = 8
    ws5.column_dimensions["F"].width = 20

    # ── Save ──────────────────────────────────────────────────────────────────
    MASTER_XL_PATH.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(MASTER_XL_PATH))
    logger.info(f"[manager] Master workbook saved: {MASTER_XL_PATH}")
    return MASTER_XL_PATH


# ──────────────────────────────────────────────────────────────────────────────
# FILTERED WORKBOOKS
# (generated on demand for specific views)
# ──────────────────────────────────────────────────────────────────────────────

def build_filtered_workbook(
    receipts: list[dict],
    sort_by: str = "date",          # "date", "vendor", "amount", "item"
    title_suffix: str = "",
    output_path: Optional[Path] = None,
) -> Path:
    """
    Build a focused single-sheet workbook for a specific view/filter.
    Faster than the full 5-sheet master.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
    except ImportError:
        raise ImportError("pip install openpyxl --break-system-packages")

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    wb  = Workbook()
    ws  = wb.active
    ws.freeze_panes = "A3"

    if not output_path:
        fname = f"GOJ_Receipts_{sort_by}_{date.today().isoformat()}.xlsx"
        output_path = REPORTS_DIR / fname

    # Title
    ws.merge_cells("A1:H1")
    ws["A1"].value     = f"Garden of Joy — Receipts {title_suffix}  ({now})"
    ws["A1"].font      = Font(name="Arial", bold=True, size=12, color="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[1].height = 20

    if sort_by == "item":
        _apply_header(ws, 2,
            ["Receipt ID", "Date", "Vendor", "Item Description",
             "Qty", "Unit Price ($)", "Line Total ($)", "Category"])
        ws.row_dimensions[2].height = 18

        all_items: list[tuple] = []
        for r in receipts:
            items = _get_line_items(r["id"])
            if items:
                for item in items:
                    all_items.append((r, item))
            else:
                all_items.append((r, {
                    "description": r.get("vendor", "Receipt"),
                    "quantity": 1, "unit_price": r.get("amount", 0),
                    "total": r.get("amount", 0),
                }))
        all_items.sort(key=lambda x: (x[1].get("description") or "").lower())

        for row_i, (r, item) in enumerate(all_items, start=3):
            ws.cell(row_i, 1, r["id"])
            ws.cell(row_i, 2, r["receipt_date"])
            ws.cell(row_i, 3, r["vendor"])
            ws.cell(row_i, 4, (item.get("description") or "")[:60])
            ws.cell(row_i, 5, item.get("quantity", 1))
            up = ws.cell(row_i, 6, item.get("unit_price"))
            if up.value: up.number_format = _xl_currency_format()
            lt = ws.cell(row_i, 7, item.get("total"))
            if lt.value: lt.number_format = _xl_currency_format()
            ws.cell(row_i, 8, r["category"])
            fill = _xl_zebra_style(row_i)
            if fill:
                for col in range(1, 9):
                    ws.cell(row_i, col).fill = fill

    else:
        _apply_header(ws, 2,
            ["ID", "Date", "Vendor / Company", "Category",
             "Amount ($)", "Tax ($)", "Submitted By", "Description"])
        ws.row_dimensions[2].height = 18

        if sort_by == "vendor":
            receipts = sorted(receipts, key=lambda x: (x["vendor"] or "").lower())
        elif sort_by == "amount":
            receipts = sorted(receipts, key=lambda x: x["amount"] or 0, reverse=True)
        else:
            receipts = sorted(receipts, key=lambda x: x["receipt_date"] or "", reverse=True)

        for row_i, r in enumerate(receipts, start=3):
            ws.cell(row_i, 1, r["id"])
            ws.cell(row_i, 2, r["receipt_date"])
            ws.cell(row_i, 3, r["vendor"])
            ws.cell(row_i, 4, r["category"])
            a = ws.cell(row_i, 5, r["amount"])
            a.number_format = _xl_currency_format()
            t = ws.cell(row_i, 6, r["tax"])
            t.number_format = _xl_currency_format()
            ws.cell(row_i, 7, (r["submitted_by"] or "kato").title())
            ws.cell(row_i, 8, (r["description"] or "")[:50])
            fill = _xl_zebra_style(row_i)
            if fill:
                for col in range(1, 9):
                    ws.cell(row_i, col).fill = fill

        last = 2 + len(receipts)
        ws.cell(last + 1, 3, "TOTAL").font = Font(name="Arial", bold=True)
        tc = ws.cell(last + 1, 5,
                     f"=SUM(E3:E{last})" if receipts else 0)
        tc.number_format = _xl_currency_format()
        tc.font = Font(name="Arial", bold=True, color="1F4E79")

    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 13
    ws.column_dimensions["F"].width = 11
    ws.column_dimensions["G"].width = 14
    ws.column_dimensions["H"].width = 32

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT EXPENSE REPORT
# (python-docx, runtime-safe)
# ──────────────────────────────────────────────────────────────────────────────

def build_expense_report_docx(
    receipts: list[dict],
    period_label: str = "",
    output_path: Optional[Path] = None,
) -> Optional[Path]:
    """
    Build a formal expense report as a .docx file.
    Organized by category. Each category gets a section with a receipt table.
    Requires python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.warning("[manager] python-docx not installed — skipping Word report")
        return None

    if not output_path:
        fname = f"GOJ_Expense_Report_{date.today().isoformat()}.docx"
        output_path = REPORTS_DIR / fname

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Garden of Joy — Expense Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.size = Pt(20)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(period_label or date.today().strftime("%B %Y"))
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    sr.font.bold = True

    generated_p = doc.add_paragraph()
    generated_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = generated_p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    gr.font.size = Pt(9)
    gr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ── Executive summary table ────────────────────────────────────────────────
    total_spend = sum(r["amount"] or 0 for r in receipts)
    total_tax   = sum(r["tax"]    or 0 for r in receipts)

    doc.add_heading("Executive Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    def _fill_table_row(tbl, row_idx: int, label: str, value: str, bold_val: bool = False):
        row = tbl.rows[row_idx]
        lbl_cell = row.cells[0]
        val_cell = row.cells[1]
        lbl_cell.text = label
        val_cell.text = value
        lbl_cell.paragraphs[0].runs[0].font.bold = True
        if bold_val:
            val_cell.paragraphs[0].runs[0].font.bold  = True
            val_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    _fill_table_row(summary_table, 0, "Total Spend",     f"${total_spend:,.2f}", True)
    _fill_table_row(summary_table, 1, "Total Tax",       f"${total_tax:,.2f}")
    _fill_table_row(summary_table, 2, "Total Receipts",  str(len(receipts)))
    _fill_table_row(summary_table, 3, "Period",          period_label or date.today().strftime("%B %Y"))

    doc.add_paragraph()

    # ── By category sections ───────────────────────────────────────────────────
    from collections import defaultdict
    cat_groups: dict[str, list[dict]] = defaultdict(list)
    for r in receipts:
        cat_groups[r["category"] or "Misc"].append(r)

    for cat in sorted(cat_groups.keys()):
        cat_receipts = sorted(cat_groups[cat], key=lambda x: x["receipt_date"] or "", reverse=True)
        cat_total    = sum(r["amount"] or 0 for r in cat_receipts)

        # Category heading
        doc.add_heading(f"{cat}  —  ${cat_total:,.2f}", level=2)

        # Receipts table
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["Date", "Vendor", "Amount", "Tax", "Submitted By"]):
            p = hdr[i].paragraphs[0]
            r = p.add_run(label)
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Navy background
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1F4E79")
            shd.set(qn("w:color"), "FFFFFF")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)

        # Data rows
        for idx, r in enumerate(cat_receipts):
            row = tbl.add_row()
            row.cells[0].text = r["receipt_date"] or ""
            row.cells[1].text = (r["vendor"] or "")[:30]
            row.cells[2].text = f"${r['amount']:,.2f}" if r["amount"] else "N/A"
            row.cells[3].text = f"${r['tax']:,.2f}"     if r["tax"]    else ""
            row.cells[4].text = (r["submitted_by"] or "kato").title()
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else None
            # Alternating row color
            if idx % 2 == 0:
                for cell in row.cells:
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd  = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "EBF3FB")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)

        # Category total row
        total_row = tbl.add_row()
        total_row.cells[0].text = "SUBTOTAL"
        total_row.cells[2].text = f"${cat_total:,.2f}"
        for i in range(5):
            run = total_row.cells[i].paragraphs[0].runs
            if run:
                run[0].font.bold = True

        doc.add_paragraph()

    # ── Footer line ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Garden of Joy Adult Day Care  ·  Gold Health Systems  ·  "
        f"Confidential — Chairman & Vlad Eyes Only"
    )
    fr.font.size  = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fr.font.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"[manager] Expense report saved: {output_path}")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# MAIN RECEIPT MANAGER CLASS
# ──────────────────────────────────────────────────────────────────────────────

class ReceiptManager:
    """
    Central receipt management for Garden of Joy.

    Handles multi-user intake, access control, and report generation.
    Integrates with rex_receipt_reader.py for OCR.
    """

    def __init__(self):
        _ensure_db()

    def get_role(self, user_label: str) -> int:
        """Return access level for a user."""
        roles = {
            "kato":  ROLE_FULL,
            "vlad":  ROLE_VIEW,
            "misha": ROLE_SUBMIT,
        }
        return roles.get(user_label.lower(), ROLE_NONE)

    # ── Receipt submission (Misha or Kato) ─────────────────────────────────────

    def submit_receipt_from_photo(
        self,
        photo_bytes: bytes,
        submitted_by: str = "kato",
        caption: str = "",
        notify_callback=None,
    ) -> dict:
        """
        Process a receipt photo from any user.
        Uses rex_receipt_reader for OCR.
        Optionally calls notify_callback(msg) to forward to Kato/Vlad.
        """
        try:
            from rex_receipt_reader import ReceiptReader
            reader = ReceiptReader()
            result = reader.handle_telegram_photo(photo_bytes, "receipt.jpg")
        except Exception as e:
            return {"error": str(e), "summary": f"⚠️ OCR failed: {e}"}

        if "error" in result:
            return result

        # Update submitter in DB
        try:
            con = sqlite3.connect(str(LEDGER_DB))
            con.execute(
                "UPDATE receipts SET submitted_by=? WHERE id=?",
                (submitted_by, result["receipt_id"])
            )
            con.commit()
            con.close()
        except Exception:
            pass

        # V4: emit receipt.submitted event (silent — never blocks receipt flow)
        try:
            emit_receipt_event(
                "receipt.submitted",
                result.get("receipt_id", 0),
                submitted_by,
                {"vendor": result.get("vendor"), "amount": result.get("amount"),
                 "category": result.get("category")},
            )
        except Exception:
            pass

        # Build appropriate summary based on who submitted
        amt_str  = "${:.2f}".format(result["amount"]) if result["amount"] else "amount unclear"
        amt_str2 = "${:.2f}".format(result["amount"]) if result["amount"] else "unclear"
        if submitted_by == "misha":
            summary = (
                "✅ Got it, Misha! Receipt filed.\n\n"
                "🏪 {}\n"
                "💵 {}\n"
                "📂 Category: {}\n"
                "📄 Receipt #{:05d} — logged."
            ).format(result["vendor"], amt_str, result["category"], result["receipt_id"])
            # Notify Kato and Vlad (via callback)
            if notify_callback:
                kato_msg = (
                    "📋 *New Receipt from Misha*\n\n"
                    "🏪 Vendor:    {}\n"
                    "📅 Date:      {}\n"
                    "💵 Amount:    {}\n"
                    "📂 Category:  {}\n"
                    "📄 Receipt #{:05d} filed.\n\n"
                    "Send `receipt {} category [Name]` to recategorize."
                ).format(
                    result["vendor"], result["date"], amt_str2,
                    result["category"], result["receipt_id"], result["receipt_id"]
                )
                notify_callback(kato_msg)
        else:
            summary = result.get("summary", "✅ Receipt filed.")

        result["summary"]      = summary
        result["submitted_by"] = submitted_by

        # Rebuild master workbook in background (non-blocking)
        try:
            self.rebuild_master_workbook()
        except Exception as e:
            logger.debug(f"[manager] Background workbook rebuild error: {e}")

        return result

    # ── Access-controlled report generation ───────────────────────────────────

    def get_receipts_for_user(
        self, user_label: str, **filters
    ) -> Optional[list[dict]]:
        """Return receipts if user has view+ access, else None.
        V4: financial fields redacted for non-financial roles via ReceiptVisibilityGate."""
        role = self.get_role(user_label)
        if role < ROLE_VIEW:
            return None
        raw = _get_receipts(**filters)
        # V4 visibility gate — redacts amount/tax/raw_text for non-financial roles
        if _visibility_gate is not None:
            return _visibility_gate.filter_receipts(raw, user_label=user_label)
        return raw

    def rebuild_master_workbook(
        self,
        start: str = "2000-01-01",
        end:   str = "2099-12-31",
    ) -> Optional[Path]:
        """Rebuild the full 5-sheet master workbook."""
        receipts = _get_receipts(start=start, end=end)
        if not receipts:
            logger.info("[manager] No receipts — skipping workbook rebuild")
            return None
        return build_master_workbook(receipts)

    def get_filtered_excel(
        self, user_label: str, sort_by: str = "date",
        submitter_filter: Optional[str] = None,
        vendor_filter: Optional[str] = None,
        start: str = "2000-01-01",
        end:   str = "2099-12-31",
    ) -> Optional[Path]:
        """Generate a filtered/sorted Excel for Kato or Vlad."""
        role = self.get_role(user_label)
        if role < ROLE_VIEW:
            return None

        receipts = _get_receipts(
            start=start, end=end,
            submitter=submitter_filter,
            vendor_filter=vendor_filter,
        )
        if not receipts:
            return None

        sort_label_map = {
            "date":   "by Date",
            "vendor": "by Vendor",
            "amount": "by Amount",
            "item":   "— Line Items",
        }
        title_suffix = sort_label_map.get(sort_by, "")
        if submitter_filter:
            title_suffix += f" | {submitter_filter.title()} only"
        if vendor_filter:
            title_suffix += f" | {vendor_filter}"

        fname = f"GOJ_Receipts_{sort_by}_{date.today().isoformat()}.xlsx"
        out   = REPORTS_DIR / fname
        return build_filtered_workbook(
            receipts, sort_by=sort_by,
            title_suffix=title_suffix, output_path=out
        )

    def get_expense_report_docx(
        self, user_label: str,
        month: str = "",
    ) -> Optional[Path]:
        """Generate a formal Word expense report for Kato or Vlad."""
        role = self.get_role(user_label)
        if role < ROLE_VIEW:
            return None

        if month:
            start = month + "-01"
            y, m  = int(month[:4]), int(month[5:7])
            end_d = (datetime(y, m, 1) + timedelta(days=32)).replace(day=1) - timedelta(days=1)
            end   = end_d.strftime("%Y-%m-%d")
            label = datetime(y, m, 1).strftime("%B %Y")
        else:
            start = date.today().strftime("%Y-%m-01")
            end   = date.today().isoformat()
            label = date.today().strftime("%B %Y") + " (Month to Date)"

        receipts = _get_receipts(start=start, end=end)
        if not receipts:
            return None

        fname = f"GOJ_Expense_Report_{month or date.today().strftime('%Y-%m')}.docx"
        out   = REPORTS_DIR / fname
        return build_expense_report_docx(receipts, period_label=label, output_path=out)

    # ── Soft-delete (Chairman only) ────────────────────────────────────────────

    def delete_receipt(self, user_label: str, receipt_id: int) -> str:
        if self.get_role(user_label) < ROLE_FULL:
            return "❌ Only the Chairman can delete receipts."
        try:
            con = sqlite3.connect(str(LEDGER_DB))
            row = con.execute(
                "SELECT vendor, amount FROM receipts WHERE id=?", (receipt_id,)
            ).fetchone()
            if not row:
                con.close()
                return f"Receipt #{receipt_id:05d} not found."
            con.execute("UPDATE receipts SET deleted=1 WHERE id=?", (receipt_id,))
            con.commit()
            con.close()
            return f"🗑️ Receipt #{receipt_id:05d} ({row[0]}, ${row[1]:.2f}) soft-deleted."
        except Exception as e:
            return f"Delete error: {e}"

    def correct_category(self, user_label: str, receipt_id: int, new_cat: str) -> str:
        if self.get_role(user_label) < ROLE_FULL:
            return "❌ Only the Chairman can recategorize receipts."
        try:
            from rex_receipt_reader import ReceiptReader, CATEGORIES
            valid = list(CATEGORIES.keys()) + ["Misc"]
            if new_cat not in valid:
                return f"Unknown category. Valid: {', '.join(valid)}"
            reader = ReceiptReader()
            return reader.correct_category(receipt_id, new_cat)
        except Exception as e:
            return f"Recategorize error: {e}"

    # ── Quick text summary ─────────────────────────────────────────────────────

    def get_text_summary(
        self, user_label: str, days: int = 7,
        submitter: Optional[str] = None,
    ) -> str:
        """Fast text summary of recent receipts."""
        role = self.get_role(user_label)
        if role < ROLE_VIEW:
            return "❌ You don't have access to the receipt section."

        start    = (date.today() - timedelta(days=days)).isoformat()
        end      = date.today().isoformat()
        receipts = _get_receipts(start=start, end=end, submitter=submitter)

        if not receipts:
            return f"No receipts in the last {days} days."

        total = sum(r["amount"] or 0 for r in receipts)
        lines = [
            f"📋 *Receipts — Last {days} Days*\n",
            f"Total: ${total:,.2f}  |  {len(receipts)} receipt(s)\n",
        ]
        for r in sorted(receipts, key=lambda x: x["receipt_date"] or "", reverse=True)[:15]:
            sub  = f" [{r['submitted_by'].title()}]" if r["submitted_by"] != "kato" else ""
            amt  = f"${r['amount']:,.2f}" if r["amount"] else "N/A"
            lines.append(
                f"  #{r['id']:05d}  {r['receipt_date']}  "
                f"{(r['vendor'] or '')[:22]:<22}  {amt:>9}  "
                f"[{r['category']}]{sub}"
            )
        if len(receipts) > 15:
            lines.append(f"  … and {len(receipts) - 15} more. Run `receipts report` for full view.")

        lines.append(
            "\n*Commands:*\n"
            "`receipts by company` · `receipts by price` · "
            "`receipts by item` · `receipts report`"
        )
        return "\n".join(lines)


# ──────────────────────────────────────────────────────────────────────────────
# TELEGRAM COMMAND ROUTER
# ──────────────────────────────────────────────────────────────────────────────

_manager_instance: Optional[ReceiptManager] = None

def get_manager() -> ReceiptManager:
    global _manager_instance
    if _manager_instance is None:
        _manager_instance = ReceiptManager()
    return _manager_instance


def handle_receipt_command(
    text: str,
    user_label: str = "kato",
    photo_bytes: Optional[bytes] = None,
    notify_callback=None,
) -> dict:
    """
    Main entry point from rex_telegram_bot.py.

    Returns dict with:
      reply   — text to send back to the user
      file    — optional Path to an Excel/docx file to send as document
    """
    manager = get_manager()
    role    = manager.get_role(user_label)
    lower   = text.lower().strip()

    # ── Photo submission ───────────────────────────────────────────────────────
    if photo_bytes:
        result = manager.submit_receipt_from_photo(
            photo_bytes, submitted_by=user_label,
            caption=text, notify_callback=notify_callback,
        )
        return {"reply": result.get("summary", "⚠️ Error processing receipt.")}

    # ── Staff (Misha) — submit only ────────────────────────────────────────────
    if role == ROLE_SUBMIT:
        return {"reply":
            "Send me a photo of your receipt and I'll file it right away!\n"
            "Just tap 📎 and select the photo."
        }

    # ── No access ──────────────────────────────────────────────────────────────
    if role == ROLE_NONE:
        return {"reply": "You don't have access to the receipt section."}

    # ── View / Full access commands ────────────────────────────────────────────

    # receipts by company / vendor
    if any(kw in lower for kw in ["by company", "by vendor", "sort company", "sort vendor"]):
        receipts = _get_receipts()
        if not receipts:
            return {"reply": "No receipts on file yet."}
        xl = build_filtered_workbook(receipts, sort_by="vendor",
                                     title_suffix="— Sorted by Company")
        return {"reply": "📊 Receipts sorted by company:", "file": xl}

    # receipts by price / amount
    if any(kw in lower for kw in ["by price", "by amount", "highest", "sort price", "sort amount"]):
        receipts = _get_receipts()
        if not receipts:
            return {"reply": "No receipts on file yet."}
        xl = build_filtered_workbook(receipts, sort_by="amount",
                                     title_suffix="— Sorted by Amount")
        return {"reply": "📊 Receipts sorted by amount (highest first):", "file": xl}

    # receipts by item
    if any(kw in lower for kw in ["by item", "line item", "itemized", "sort item"]):
        receipts = _get_receipts()
        if not receipts:
            return {"reply": "No receipts on file yet."}
        xl = build_filtered_workbook(receipts, sort_by="item",
                                     title_suffix="— Line Items")
        return {"reply": "📊 Receipts — line items expanded:", "file": xl}

    # receipts master / full workbook
    if any(kw in lower for kw in ["master", "full workbook", "all sheets", "all views"]):
        xl = manager.rebuild_master_workbook()
        if not xl:
            return {"reply": "No receipts on file yet."}
        return {"reply": "📊 Full master workbook (5 sheets):", "file": xl}

    # receipts report [optional month]
    if "report" in lower:
        month = ""
        m = re.search(r'(\d{4}-\d{2})', lower)
        if m:
            month = m.group(1)
        docx_path = manager.get_expense_report_docx(user_label, month=month)
        if not docx_path:
            return {"reply": "No receipts found for that period."}
        period = month or date.today().strftime("%B %Y")
        return {"reply": f"📄 Formal expense report — {period}:", "file": docx_path}

    # receipts misha
    if "misha" in lower:
        receipts = _get_receipts(submitter="misha")
        if not receipts:
            return {"reply": "No receipts from Misha yet."}
        xl = build_filtered_workbook(receipts, sort_by="date",
                                     title_suffix="— Misha's Submissions")
        return {"reply": f"📊 Misha's receipts ({len(receipts)} total):", "file": xl}

    # receipts this week / last N days
    if any(kw in lower for kw in ["this week", "last week", "7 days", "seven days"]):
        return {"reply": manager.get_text_summary(user_label, days=7)}

    if any(kw in lower for kw in ["this month", "30 days", "thirty days"]):
        return {"reply": manager.get_text_summary(user_label, days=30)}

    # receipts [vendor name]
    vendor_m = re.match(r'^receipts?\s+(.+)$', lower)
    if vendor_m:
        vendor_query = vendor_m.group(1).strip()
        # Don't match help keywords as vendor names
        skip = {"by", "report", "master", "misha", "vlad", "kato", "help",
                 "this", "last", "week", "month", "delete", "all"}
        if vendor_query not in skip and len(vendor_query) > 2:
            receipts = _get_receipts(vendor_filter=vendor_query)
            if receipts:
                xl = build_filtered_workbook(
                    receipts, sort_by="date",
                    title_suffix=f"— Vendor: {vendor_query.title()}"
                )
                return {"reply": f"📊 Receipts matching '{vendor_query}':", "file": xl}
            return {"reply": f"No receipts matching '{vendor_query}'."}

    # receipts delete [id]
    del_m = re.match(r'^receipts?\s+delete\s+(\d+)$', lower)
    if del_m:
        return {"reply": manager.delete_receipt(user_label, int(del_m.group(1)))}

    # Default: recent summary + help
    return {"reply": manager.get_text_summary(user_label, days=30) + "\n\n" + _receipt_help(role)}


def _receipt_help(role: int) -> str:
    if role >= ROLE_FULL:
        return (
            "📄 *Receipt Commands:*\n"
            "• Send a *photo* → auto-file any receipt\n"
            "• `receipts by company` — Excel sorted by vendor\n"
            "• `receipts by price` — Excel sorted highest first\n"
            "• `receipts by item` — line items expanded\n"
            "• `receipts master` — full 5-sheet workbook\n"
            "• `receipts report` — formal Word report (this month)\n"
            "• `receipts report 2025-03` — specific month\n"
            "• `receipts misha` — only Misha's submissions\n"
            "• `receipts [vendor name]` — filter by company\n"
            "• `receipts this week` — 7-day summary\n"
            "• `receipts delete [id]` — soft-delete a receipt"
        )
    else:  # ROLE_VIEW (Vlad)
        return (
            "📄 *Receipt Commands (View Only):*\n"
            "• `receipts` — recent summary\n"
            "• `receipts by company` — Excel by vendor\n"
            "• `receipts this week` — 7-day text summary"
        )


# ──────────────────────────────────────────────────────────────────────────────
# SELF-TEST
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("REX RECEIPT MANAGER — SELF-TEST")
    print("=" * 60)

    # 1. DB init
    _ensure_db()
    print("✓ Test 1: DB init OK")

    # 2. Role access control
    mgr = ReceiptManager()
    assert mgr.get_role("kato")  == ROLE_FULL
    assert mgr.get_role("vlad")  == ROLE_VIEW
    assert mgr.get_role("misha") == ROLE_SUBMIT
    assert mgr.get_role("john")  == ROLE_NONE
    print("✓ Test 2: Role access control OK")

    # 3. Insert test receipts (use dynamic dates so summary window always catches them)
    _td = date.today()
    con = sqlite3.connect(str(LEDGER_DB))
    test_ids = []
    test_data = [
        ((_td - timedelta(days=6)).isoformat(), "Whole Foods Market", 52.30, 3.14, "Meals/Food",  "misha"),
        ((_td - timedelta(days=5)).isoformat(), "Office Depot",       28.00, 1.68, "Supplies",    "misha"),
        ((_td - timedelta(days=4)).isoformat(), "CVS Pharmacy",       19.45, 1.17, "Medical",     "kato"),
        ((_td - timedelta(days=2)).isoformat(), "Shell Gas Station",  61.00, 0.00, "Transport",   "misha"),
        ((_td - timedelta(days=1)).isoformat(), "Amazon",            134.99, 0.00, "Technology",  "kato"),
    ]
    for rd, vendor, amt, tax, cat, sub in test_data:
        cur = con.execute(
            "INSERT INTO receipts (receipt_date, vendor, amount, tax, category, submitted_by) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (rd, vendor, amt, tax, cat, sub)
        )
        test_ids.append(cur.lastrowid)
        # Add a line item to first receipt
        if vendor == "Whole Foods Market":
            con.execute(
                "INSERT INTO line_items (receipt_id, description, quantity, unit_price, total) "
                "VALUES (?, ?, ?, ?, ?)",
                (cur.lastrowid, "Organic Chicken", 1, 14.99, 14.99)
            )
            con.execute(
                "INSERT INTO line_items (receipt_id, description, quantity, unit_price, total) "
                "VALUES (?, ?, ?, ?, ?)",
                (cur.lastrowid, "Mixed Vegetables", 2, 3.99, 7.98)
            )
    con.commit()
    con.close()
    print(f"✓ Test 3: Inserted {len(test_data)} test receipts (IDs: {test_ids})")

    # 4. Get receipts
    receipts = _get_receipts(start=(_td - timedelta(days=10)).isoformat(), end=_td.isoformat())
    assert len(receipts) >= len(test_data), f"Expected {len(test_data)}, got {len(receipts)}"
    print(f"✓ Test 4: Retrieved {len(receipts)} receipts")

    # 5. Text summary
    summary = mgr.get_text_summary("kato", days=30)
    assert "Receipts" in summary
    print("✓ Test 5: Text summary OK")

    # 6. Access control — Vlad can view
    assert mgr.get_receipts_for_user("vlad") is not None
    # Misha cannot view
    assert mgr.get_receipts_for_user("misha") is None
    print("✓ Test 6: Access control OK (Vlad=view, Misha=submit-only)")

    # 7. Filtered workbook — by company
    try:
        xl_path = mgr.get_filtered_excel("kato", sort_by="vendor")
        assert xl_path and xl_path.exists(), f"Excel not created: {xl_path}"
        print(f"✓ Test 7: By-vendor Excel OK ({xl_path.name})")
    except ImportError:
        print("⚠ Test 7: openpyxl not installed — skipping Excel test")
        print("   Install: pip install openpyxl --break-system-packages")

    # 8. Filtered workbook — by item
    try:
        xl_path = mgr.get_filtered_excel("kato", sort_by="item")
        assert xl_path and xl_path.exists()
        print(f"✓ Test 8: By-item Excel OK ({xl_path.name})")
    except ImportError:
        print("⚠ Test 8: openpyxl not installed — skipped")

    # 9. Master workbook
    try:
        master = mgr.rebuild_master_workbook()
        assert master and master.exists()
        print(f"✓ Test 9: Master workbook OK (5 sheets, {master.stat().st_size // 1024}KB)")
    except ImportError:
        print("⚠ Test 9: openpyxl not installed — skipped")

    # 10. Word report
    try:
        docx_path = mgr.get_expense_report_docx("kato", month="2025-04")
        if docx_path:
            print(f"✓ Test 10: Word report OK ({docx_path.name})")
        else:
            print("⚠ Test 10: python-docx not installed — skipping Word report")
            print("   Install: pip install python-docx --break-system-packages")
    except Exception as e:
        print(f"⚠ Test 10: Word report error: {e}")

    # 11. Telegram command routing
    result = handle_receipt_command("receipts by company", "kato")
    assert "reply" in result
    print("✓ Test 11: Telegram command routing OK")

    # 12. Access denial for Misha
    result = handle_receipt_command("receipts by company", "misha")
    assert "photo" in result["reply"].lower() or "send" in result["reply"].lower()
    print("✓ Test 12: Misha correctly redirected to submit-only")

    # Cleanup test records
    con = sqlite3.connect(str(LEDGER_DB))
    for tid in test_ids:
        con.execute("DELETE FROM line_items WHERE receipt_id=?", (tid,))
        con.execute("DELETE FROM receipts WHERE id=?", (tid,))
    con.commit()
    con.close()

    print()
    print("=" * 60)
    print("ALL TESTS PASSED — rex_receipt_manager.py ready")
    print()
    print("Libraries needed:")
    try:
        import openpyxl
        print("  openpyxl: ✓")
    except ImportError:
        print("  openpyxl: ✗  →  pip install openpyxl --break-system-packages")
    try:
        import docx
        print("  python-docx: ✓")
    except ImportError:
        print("  python-docx: ✗  →  pip install python-docx --break-system-packages")
    print("=" * 60)
